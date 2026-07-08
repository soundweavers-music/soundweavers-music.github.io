"""
世界樂器百科電子書輸出 — 依 A1~A9 分類輸出 5 本 .docx，再轉 .pdf

每本書：
  1. 封面頁（設計排版）
  2. 關於作者
  3. 目錄頁（Word TOC 欄位，自動對應頁碼與導覽）
  4. 該分類所有樂器完整介紹（每樂器獨立分頁，YouTube QR Code 橫排）

書名：世界樂器百科Ⅰ ~ Ⅴ
Output: outputs/world-instruments-static/assets/
"""

import os
import re
import subprocess
import sys
from collections import defaultdict
from pathlib import Path
from io import BytesIO

import qrcode
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = Path(__file__).resolve().parent.parent
CONTENT_DIR = BASE_DIR / "content" / "instruments"
OUTPUT_DIR = BASE_DIR / "outputs" / "world-instruments-static" / "assets"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

ROMAN = {1: "Ⅰ", 2: "Ⅱ", 3: "Ⅲ", 4: "Ⅳ", 5: "Ⅴ"}

CATEGORY_ARCHITECTURE = {
    "A1": {"name": "吹奏與氣息樂器", "subtitle": "靠氣流、管身、簧片、唇振、風箱或風管系統發聲", "roman": "Ⅰ", "filename": "A1-吹奏與氣息樂器"},
    "A2": {"name": "弦樂器", "subtitle": "靠弦震動發聲；撥弦、擦弦、擊弦到鍵控弦鳴", "roman": "Ⅱ", "filename": "A2-弦樂器"},
    "A3": {"name": "鼓與打擊樂器", "subtitle": "膜鳴與體鳴；鼓、鑼鐘、木石竹到舌片手碟", "roman": "Ⅲ", "filename": "A3-鼓與打擊樂器"},
    "A4": {"name": "電子與電聲樂器", "subtitle": "電子振盪、合成、取樣、電聲改造到數位控制", "roman": "Ⅳ", "filename": "A4-電子與電聲樂器"},
    "A9": {"name": "待分類／偵錯暫存", "subtitle": "名稱待查、發聲原理待查、重複合併、來源不足", "roman": "Ⅴ", "filename": "A9-待分類"},
}

ABOUT_TEXT = (
    "世界聲音百科由「隔壁織音人」團隊整理，致力於建立一個兼具知識性、"
    "可讀性與探索感的世界樂器百科平台。我們將世界各地的樂器整理為系統化"
    "的知識體系，讓讀者能像翻閱地圖般，循著聲音的軌跡，走進不同文化的現場。\n\n"
    "作者：隔壁織音人 (Next Door Sound Weavers)\n"
    "網站：https://soundweavers-music.github.io/\n"
    "YouTube：https://www.youtube.com/@NextDoorSoundWeavers/\n"
    "電子郵件：nextdoor.soundweavers@gmail.com\n\n"
    "服務項目：編曲製作、詞曲訂製、後期處理、人聲錄製、混音母帶。"
)


def parse_frontmatter(text):
    m = re.match(r'^---\n(.*?)\n---', text, re.DOTALL)
    if not m:
        return {}, text.strip()
    data = {}
    for line in m.group(1).split('\n'):
        if ':' not in line:
            continue
        k, v = line.split(':', 1)
        data[k.strip()] = v.strip()
    return data, text[m.end():].strip()


def make_qr_buf(url):
    if not url:
        return None
    qr = qrcode.QRCode(box_size=2, border=1)
    qr.add_data(url)
    qr.make(fit=True)
    buf = BytesIO()
    qr.make_image(fill_color="black", back_color="white").save(buf, format="PNG")
    buf.seek(0)
    return buf


def read_all_instruments():
    instruments = []
    for path in sorted(CONTENT_DIR.glob("*.md")):
        meta, body = parse_frontmatter(path.read_text(encoding="utf-8"))
        body = body.strip()
        sections = {}
        for head, key in {"介紹": "introduction", "歷史背景": "history", "音色描述": "timbre",
                          "樂器材質": "material", "教學": "tutorial"}.items():
            m = re.search(rf"^##\s*{re.escape(head)}\s*\n(.*?)(?=\n## |\Z)", body, re.DOTALL | re.MULTILINE)
            if m:
                sections[key] = m.group(1).strip()
        instruments.append({
            "slug": path.stem,
            "class_code": meta.get("class_code", ""),
            "frontend_class": meta.get("frontend_class", ""),
            "subcategory": meta.get("subcategory", ""),
            "title_zh": meta.get("title_zh", path.stem),
            "title_original": meta.get("title_original", ""),
            "family_std": meta.get("family_std", ""),
            "sound_hs": meta.get("sound_hs", ""),
            "playing_method": meta.get("playing_method", ""),
            "interface_tags": meta.get("interface_tags", ""),
            "region_culture": meta.get("region_culture", ""),
            "listening_sound_tags": meta.get("listening_sound_tags", ""),
            "ensemble_links": meta.get("ensemble_links", ""),
            "verification_status": meta.get("verification_status", ""),
            "issue_note": meta.get("issue_note", ""),
            "source_url": meta.get("source_url", ""),
            "image": meta.get("image", ""),
            "youtube_ids": meta.get("youtube_ids", ""),
            "introduction": sections.get("introduction", ""),
            "history": sections.get("history", ""),
            "timbre": sections.get("timbre", ""),
            "material": sections.get("material", ""),
            "tutorial": sections.get("tutorial", ""),
        })
    return instruments


def find_font(name_substrings):
    search_dirs = ["C:/Windows/Fonts", "/usr/share/fonts", "/usr/local/share/fonts"]
    exts = (".ttc", ".ttf", ".otf")
    for sd in search_dirs:
        if not os.path.isdir(sd):
            continue
        for root, _, files in os.walk(sd):
            for fn in files:
                if any(kw.lower() in fn.lower() for kw in name_substrings) and fn.lower().endswith(exts):
                    return os.path.join(root, fn)
    return None


FONT_SERIF = find_font(["simsun", "songti", "noto serif cjk", "notoserifcjk", "droid serif"])
FONT_SANS = find_font(["msjh", "jhenghei", "noto sans cjk", "notosanscjk", "wqy"])

print(f"  [字型] 襯線: {FONT_SERIF or '未找到'}")
print(f"  [字型] 無襯線: {FONT_SANS or '未找到'}")


def set_run_font(run, name_serif=None, name_sans=None, size=None, bold=None, color=None, italic=None):
    if name_serif:
        run.font.name = name_serif
        rpr = run._r.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rpr.insert(0, rf)
        rf.set(qn("w:eastAsia"), name_serif)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if italic is not None:
        run.font.italic = italic


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run._r.append(parse_xml(f'<w:br {nsdecls("w")} type="page"/>'))
    return p


def set_spacing(p, before=0, after=0, line=1.5):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def clean_md(text):
    return re.sub(r"[#*_~`>\[\]()|]", "", text)


def ensure_section_break(doc):
    """確保文件至少有一個分節，用來設定頁尾頁碼"""
    if not doc.sections:
        doc.add_section()


def setup_multilevel_heading(doc):
    """設定多層次標題編號：第一章、1.1、1.1.1 等"""
    # 用 overrides.xml 或直接操作 numbering 太複雜，
    # python-docx 不直接支援，改用 Word 內建 heading 樣式
    pass


def build_docx(instruments, code, arch):
    """建立 DOCX 主函數"""
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

    # Normal 樣式
    style = doc.styles["Normal"]
    style.font.name = FONT_SERIF or FONT_SANS or "Microsoft JhengHei"
    style.font.size = Pt(11)
    sp = style.paragraph_format
    sp.line_spacing = 1.6
    sp.space_after = Pt(6)

    # Heading 樣式
    for lvl, sz in [(1, 22), (2, 16), (3, 13)]:
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = FONT_SANS or "Microsoft JhengHei"
        hs.font.size = Pt(sz)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════════════
    #  封面
    # ═══════════════════════════════════════════════
    for _ in range(8):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("世界樂器百科"), name_sans=FONT_SANS, size=42, bold=True, color=(0x0D, 0x76, 0x6B))
    set_spacing(p, after=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(f"世界樂器百科{arch['roman']}"), name_sans=FONT_SANS, size=20, bold=True, color=(0x34, 0x40, 0x54))
    set_spacing(p, after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(f"{code}  {arch['name']}"), name_sans=FONT_SANS, size=26, bold=True, color=(0x1A, 0x23, 0x32))
    set_spacing(p, after=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(arch['subtitle']), name_serif=FONT_SERIF, size=12, color=(0x66, 0x70, 0x85), italic=True)
    set_spacing(p, after=20)

    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(f"共 {len(instruments)} 件樂器"), name_sans=FONT_SANS, size=14, color=(0x34, 0x40, 0x54))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("隔壁織音人 · 世界聲音百科"), name_sans=FONT_SANS, size=11, color=(0x66, 0x70, 0x85))

    add_page_break(doc)

    # ═══════════════════════════════════════════════
    #  關於作者
    # ═══════════════════════════════════════════════
    doc.add_heading("關於作者", level=1)
    for line in ABOUT_TEXT.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph(line)
        set_run_font(p.runs[0], name_serif=FONT_SERIF, size=11)
        set_spacing(p, before=2, after=2)

    add_page_break(doc)

    # ═══════════════════════════════════════════════
    #  目錄（使用 Word TOC 欄位）
    # ═══════════════════════════════════════════════
    doc.add_heading("目錄", level=1)

    # 子分類列表（手寫，方便快速瀏覽）
    p = doc.add_paragraph()
    set_run_font(p.add_run(f"{code}  {arch['name']}"), name_sans=FONT_SANS, size=13, bold=True)
    set_spacing(p, before=6, after=4)

    entry_num = 1
    instruments_sorted = sorted(instruments, key=lambda x: (x.get("subcategory") or "", x.get("title_zh") or ""))
    current_sc = None
    for inst in instruments_sorted:
        sc = inst.get("subcategory") or ""
        if sc and sc != current_sc:
            current_sc = sc
            p = doc.add_paragraph()
            set_run_font(p.add_run(f"  {sc}"), name_serif=FONT_SERIF, size=11, bold=True, color=(0x34, 0x40, 0x54))
            set_spacing(p, before=4, after=1)

        title = inst["title_zh"] or inst["slug"]
        p = doc.add_paragraph()
        set_run_font(p.add_run(f"    {entry_num}. {title}"), name_serif=FONT_SERIF, size=10)
        set_spacing(p, before=1, after=1)
        entry_num += 1

    # 插入 TOC 欄位（Word 會自動抓 Heading 1~3 的標題填入頁碼）
    p = doc.add_paragraph()
    set_spacing(p, before=10, after=4)
    # 說明文字
    note = p.add_run("（在 Word 中按 Ctrl+A → F9 更新頁碼，或右鍵點此區域選「更新功能變數」）")
    set_run_font(note, size=9, color=(0x99, 0x99, 0x99))

    p2 = doc.add_paragraph()
    run = p2.add_run()
    run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run2 = p2.add_run()
    run2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'))
    run3 = p2.add_run()
    run3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    run4 = p2.add_run("（更新後此處將顯示各樂器的頁碼）")
    set_run_font(run4, size=9, color=(0x99, 0x99, 0x99))
    run5 = p2.add_run()
    run5._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    #  樂器內容（每樂器分頁，Heading 3 供導覽）
    # ═══════════════════════════════════════════════
    entry_num = 1
    current_sc = None
    first = True
    for inst in instruments_sorted:
        sc = inst.get("subcategory") or ""

        # 子分類標題
        if sc and sc != current_sc:
            current_sc = sc
            if not first:
                add_page_break(doc)
            doc.add_heading(sc, level=2)
        elif not first:
            add_page_break(doc)
        first = False

        # ─── 樂器名稱（Heading 3，供 TOC 與導覽） ───
        heading_text = f"{entry_num}. {inst['title_zh']}"
        doc.add_heading(heading_text, level=3)

        if inst.get("title_original"):
            p = doc.add_paragraph()
            set_run_font(p.add_run(inst["title_original"]), name_serif=FONT_SERIF, size=10, italic=True, color=(0x66, 0x70, 0x85))
            set_spacing(p, after=6)

        # ─── 欄位表格 ───
        fields = [(k, v) for k, v in [
            ("主分類代碼", inst["class_code"]), ("前台主分類", inst["frontend_class"]),
            ("子分類", inst["subcategory"]), ("樂器家族", inst["family_std"]),
            ("發聲原理／H-S", inst["sound_hs"]), ("演奏方式", inst["playing_method"]),
            ("操作介面", inst["interface_tags"]), ("地域／文化", inst["region_culture"]),
            ("聆聽與聲音標籤", inst["listening_sound_tags"]), ("常見合奏／編制", inst["ensemble_links"]),
            ("查證狀態", inst["verification_status"]), ("來源網址", inst["source_url"]),
        ] if v and v.strip()]
        if fields:
            table = doc.add_table(rows=len(fields), cols=2)
            table.style = "Light Grid Accent 1"
            for i, (k, v) in enumerate(fields):
                c0 = table.cell(i, 0)
                c0.text = ""
                set_run_font(c0.paragraphs[0].add_run(k), name_sans=FONT_SANS, size=9, bold=True)
                c1 = table.cell(i, 1)
                c1.text = ""
                set_run_font(c1.paragraphs[0].add_run(v if len(v) < 250 else v[:250] + "..."), name_serif=FONT_SERIF, size=9)
            doc.add_paragraph("")

        # ─── 內文 ───
        for title, content in [("介紹", inst["introduction"]), ("歷史背景", inst["history"]),
                               ("音色描述", inst["timbre"]), ("樂器材質", inst["material"]), ("教學", inst["tutorial"])]:
            if not content:
                continue
            p = doc.add_paragraph()
            set_run_font(p.add_run(f"【{title}】"), name_sans=FONT_SANS, size=11, bold=True, color=(0x0D, 0x76, 0x6B))
            set_spacing(p, before=6, after=2)
            clean = clean_md(content)
            for line in clean.split("\n"):
                line = line.strip()
                if not line:
                    continue
                p = doc.add_paragraph(line)
                set_run_font(p.runs[0], name_serif=FONT_SERIF, size=10)
                set_spacing(p, before=1, after=1)

        # ─── YouTube QR Code（橫排表格） ───
        yt_ids = inst.get("youtube_ids", "")
        if yt_ids:
            ids = [v.strip() for v in re.split(r"[\s,]+", yt_ids) if v.strip()][:4]
            if ids:
                # 用表格橫排呈現
                qr_table = doc.add_table(rows=2, cols=len(ids))
                qr_table.style = "Light Grid Accent 1"
                for ci, vid in enumerate(ids):
                    # 第一列：QR Code
                    qr_buf = make_qr_buf(f"https://www.youtube.com/watch?v={vid}")
                    if qr_buf:
                        cell = qr_table.cell(0, ci)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(qr_buf, width=Inches(0.35))
                    # 第二列：影片 ID
                    cell2 = qr_table.cell(1, ci)
                    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_run_font(cell2.paragraphs[0].add_run(vid), size=7, color=(0x1D, 0x4E, 0xD8))
                doc.add_paragraph("")

        entry_num += 1

    # ═══ 頁碼 ═══
    ensure_section_break(doc)
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run()
        r._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        r2 = fp.add_run()
        r2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        r3 = fp.add_run()
        r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    return doc


def convert_docx_to_pdf(docx_path, pdf_path):
    """用 LibreOffice 將 DOCX 轉為 PDF，再壓縮"""
    lo_paths = [
        "C:/Program Files/LibreOffice/program/soffice.exe",
        "C:/Program Files (x86)/LibreOffice/program/soffice.exe",
        "/usr/bin/libreoffice", "/usr/bin/soffice",
    ]
    lo_bin = None
    for p in lo_paths:
        if os.path.exists(p):
            lo_bin = p
            break
    if not lo_bin:
        import shutil
        lo_bin = shutil.which("libreoffice") or shutil.which("soffice")
    if not lo_bin:
        return False

    # 第一步：LibreOffice 轉 PDF（先用 CompressionMode=2 壓縮）
    try:
        result = subprocess.run(
            [lo_bin, "--headless", "--convert-to",
             "pdf:writer_pdf_Export:{SelectPdfVersion:=1;UseLosslessCompression:=false;Quality:=30;ReduceImageResolution:=true;MaxImageResolution:=120}",
             "--outdir", str(pdf_path.parent), str(docx_path)],
            capture_output=True, text=True, timeout=180,
        )
    except Exception as e:
        print(f"    LibreOffice 轉換失敗: {e}")
        return False

    # 第二步：用 pypdf 壓縮
    try:
        from pypdf import PdfReader, PdfWriter
        reader = PdfReader(str(pdf_path))
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        # 壓縮設定
        writer.add_metadata(reader.metadata)
        # 用較弱的壓縮保留可讀性
        with open(str(pdf_path), "wb") as f:
            writer.write(f)
        return True
    except ImportError:
        # pypdf 未安裝，直接回傳
        return result.returncode == 0
    except Exception:
        return result.returncode == 0


# ===================================================================
#  主流程
# ===================================================================
def main():
    print("讀取樂器資料...")
    all_instruments = read_all_instruments()
    print(f"  共 {len(all_instruments)} 件樂器")

    code_groups = defaultdict(list)
    for inst in all_instruments:
        code_groups[inst["class_code"] or "A9"].append(inst)

    # 檢查 LibreOffice 是否可用
    lo_available = any(os.path.exists(p) for p in [
        "C:/Program Files/LibreOffice/program/soffice.exe",
        "C:/Program Files (x86)/LibreOffice/program/soffice.exe",
        "/usr/bin/libreoffice", "/usr/bin/soffice",
    ])
    if not lo_available:
        try:
            import shutil
            lo_available = shutil.which("libreoffice") or shutil.which("soffice") is not None
        except Exception:
            pass

    for idx, (code, arch) in enumerate(CATEGORY_ARCHITECTURE.items()):
        instruments = code_groups.get(code, [])
        if not instruments:
            print(f"\n{code} {arch['name']}: 無樂器資料")
            continue

        instruments.sort(key=lambda x: (x["subcategory"] or "", x["title_zh"] or ""))
        print(f"\n{code} {arch['name']} (世界樂器百科{arch['roman']}): {len(instruments)} 件樂器")

        docx_path = OUTPUT_DIR / f"世界樂器百科{arch['roman']}_{arch['filename']}.docx"
        pdf_path = OUTPUT_DIR / f"世界樂器百科{arch['roman']}_{arch['filename']}.pdf"

        # DOCX
        try:
            doc = build_docx(instruments, code, arch)
            doc.save(str(docx_path))
            size_kb = os.path.getsize(docx_path) / 1024
            print(f"  [OK] DOCX: {docx_path.name} ({size_kb:.0f} KB)")
        except Exception as e:
            print(f"  [FAIL] DOCX: {e}")
            import traceback; traceback.print_exc()
            continue

        # DOCX → PDF
        if lo_available:
            try:
                ok = convert_docx_to_pdf(docx_path, pdf_path)
                if ok and pdf_path.exists():
                    sz = os.path.getsize(pdf_path) / 1024
                    print(f"  [OK] PDF:  {pdf_path.name} ({sz:.0f} KB)")
                else:
                    print(f"  [WARN] PDF 轉換可能失敗，請檢查 {pdf_path}")
            except Exception as e:
                print(f"  [FAIL] PDF: {e}")
        else:
            print(f"  [SKIP] PDF: 未安裝 LibreOffice，略過 PDF 轉換")
            print(f"         請自行將 DOCX 用 Word 或 LibreOffice 另存為 PDF")

    print(f"\n[OK] 全部完成!")
    print(f"   輸出資料夾: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
